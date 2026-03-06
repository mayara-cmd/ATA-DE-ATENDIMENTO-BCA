import re
import io
import os
import pandas as pd
import streamlit as st
from datetime import date
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import google.generativeai as genai

# ─────────────────────────────────────────────────────────────
# CONFIG
# ─────────────────────────────────────────────────────────────
GEMINI_API_KEY  = "AIzaSyARTNdAf6-DSBJpKKRZxQEv_GBSvzZECOk"
ORDEM_AREAS     = ["Cível", "Público", "Trabalhista", "Privado", "Compliance"]
DATA_CORTE      = pd.Timestamp("2026-12-31")

MAPA_ACAO_DEPT = {
    "reclamação trabalhista":              "Trabalhista",
    "reclamatória trabalhista":            "Trabalhista",
    "processo administrativo trabalhista": "Trabalhista",
    "mandado de segurança":                "Público",
    "execução fiscal":                     "Público",
    "processo administrativo":             "Público",
    "auditoria fiscal":                    "Público",
    "ação ordinária":                      "Público",
    "reclamação procon":                   "Cível",
    "notificação":                         "Cível",
    "indenizatória":                       "Cível",
    "ajuizamento pendente":                "Cível",
    "usucapião":                           "Cível",
    "execução de título extrajudicial":    "Cível",
    "demarcatória":                        "Cível",
    "despejo":                             "Cível",
    "falência":                            "Cível",
}

COR_TITULO = RGBColor(0x1F, 0x49, 0x7D)
COR_SUBTIT = RGBColor(0x2E, 0x74, 0xB5)

# ─────────────────────────────────────────────────────────────
# PROCESSAMENTO DO EXCEL
# ─────────────────────────────────────────────────────────────
def inferir_dept(acao, orgao):
    a, o = acao.lower().strip(), orgao.lower().strip()
    for chave, dept in MAPA_ACAO_DEPT.items():
        if chave in a:
            return dept
    if any(k in o for k in ["trabalho", "trt", "ministério público do trabalho"]):
        return "Trabalhista"
    if any(k in o for k in ["federal", "receita", "fazenda", "trf", "carf", "anp"]):
        return "Público"
    return "Cível"

def parsear_pasta(pasta):
    m = re.match(r"^((?:Proc|Serv|Doc)\s*-\s*\d+)\s*-\s*(.+)$", pasta.strip(), re.IGNORECASE)
    if m:
        return m.group(1).replace(" ", "").replace("--", "-"), m.group(2).strip()
    return pasta[:30], pasta

def carregar_excel(uploaded_file):
    df_raw = pd.read_excel(uploaded_file, dtype=str).fillna("")
    df_raw.columns = [c.strip() for c in df_raw.columns]
    df_raw = df_raw.rename(columns={
        "Pasta - Título":          "pasta",
        "Andamentos / Data/hora":  "data_str",
        "Andamentos / Descrição":  "descricao",
        "Ação":                    "acao",
        "partes":                  "partes",
        "Responsável":             "responsavel",
        "Número de CNJ":           "cnj",
        "Valor da causa":          "valor",
        "Órgão":                   "orgao",
        "Departamento":            "dept_col",
    })
    for col in ["responsavel", "cnj", "valor", "orgao", "dept_col"]:
        if col not in df_raw.columns:
            df_raw[col] = ""

    df_raw["data_parsed"] = pd.to_datetime(df_raw["data_str"], errors="coerce")
    reais = df_raw[df_raw["data_parsed"] <= DATA_CORTE].copy()

    casos = []
    for pasta, grupo in df_raw.groupby("pasta", sort=False):
        grupo_real = reais[reais["pasta"] == pasta].sort_values("data_parsed")
        primeira   = grupo.iloc[0]

        historico = "\n".join(
            f"[{str(r['data_parsed'])[:10]}] {str(r['descricao']).strip()}"
            for _, r in grupo_real.iterrows()
            if str(r["descricao"]).strip()
        )
        ultimo = grupo_real.iloc[-1]["descricao"].strip()[:300] if len(grupo_real) else ""

        dept = primeira["dept_col"].strip() if primeira["dept_col"].strip() \
               else inferir_dept(str(primeira["acao"]), str(primeira["orgao"]))

        id_caso, titulo = parsear_pasta(pasta)
        casos.append({
            "id_caso":   id_caso,
            "titulo":    titulo,
            "dept":      dept,
            "acao":      str(primeira["acao"]).strip(),
            "partes":    str(primeira["partes"]).strip(),
            "orgao":     str(primeira["orgao"]).strip(),
            "valor":     str(primeira["valor"]).strip(),
            "historico": historico,
            "ultimo":    ultimo,
            "n_and":     len(grupo_real),
        })
    return pd.DataFrame(casos)

def detectar_grupos(df_area):
    def chave(row):
        a = str(row["acao"]).lower().strip()
        p = str(row["partes"]).lower().split(" - ")[0][:40]
        u = str(row["ultimo"]).lower()[:50]
        return f"{a}|{p}|{u}"
    df2 = df_area.copy()
    df2["_k"] = df2.apply(chave, axis=1)
    return [list(g["id_caso"]) for _, g in df2.groupby("_k") if len(g) >= 2]

# ─────────────────────────────────────────────────────────────
# IA — GEMINI
# ─────────────────────────────────────────────────────────────
@st.cache_resource
def get_model():
    genai.configure(api_key=GEMINI_API_KEY)
    return genai.GenerativeModel("gemini-1.5-flash")

def resumir_caso(model, caso):
    prompt = f"""Você é assistente jurídico. Redija o resumo deste caso para ata mensal.
REGRAS:
- Máximo 5 linhas
- Comece com "Atualmente verificamos que..."
- Termine: "O próximo passo é: **[ação concreta]**."
- Foque no status atual e próxima ação. Não repita número do processo nem partes.

Tipo: {caso['acao']}
Partes: {caso['partes']}
Órgão: {caso['orgao']}
Valor: {caso['valor']}
Histórico:
{caso['historico'][-2500:]}"""
    try:
        return model.generate_content(prompt).text.strip()
    except Exception as e:
        return f"Atualmente verificamos que {caso['ultimo'][:250]}\n\nO próximo passo é: **verificar andamentos pendentes.**"

def resumir_grupo(model, ids, df):
    rows = df[df["id_caso"].isin(ids)]
    difs = "\n".join(
        f"- {r['id_caso']} ({r['partes'][:50]}): {r['ultimo'][:120]}"
        for _, r in rows.iterrows()
    )
    prompt = f"""Você é assistente jurídico. Esses casos são semelhantes — escreva UM parágrafo unificado para a ata.
REGRAS: destaque o que há de comum, mencione diferenças relevantes, termine com próximo passo em **negrito**. Máx 6 linhas.

Tipo: {rows.iloc[0]['acao']}
Casos:
{difs}"""
    try:
        return model.generate_content(prompt).text.strip()
    except:
        return f"Casos semelhantes ({', '.join(ids)}): {rows.iloc[0]['ultimo'][:200]}"

# ─────────────────────────────────────────────────────────────
# GERAÇÃO DO WORD
# ─────────────────────────────────────────────────────────────
def _shd(cell, cor):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    s = OxmlElement("w:shd")
    s.set(qn("w:val"), "clear"); s.set(qn("w:color"), "auto"); s.set(qn("w:fill"), cor)
    tcPr.append(s)

def _brd(cell):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcB = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0"); b.set(qn("w:color"), "BFBFBF")
        tcB.append(b)
    tcPr.append(tcB)

def _resumo_para(paragraph, texto):
    paragraph.clear()
    for parte in re.split(r"(\*\*.*?\*\*)", texto):
        run = paragraph.add_run(parte[2:-2] if parte.startswith("**") else parte)
        run.bold = parte.startswith("**")
        run.font.size = Pt(9)

def _subtitulo(doc, texto):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(texto)
    run.bold = True; run.font.size = Pt(11); run.font.color.rgb = COR_SUBTIT

def tabela_casos(doc, linhas):
    if not linhas:
        return
    t = doc.add_table(rows=1, cols=3)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for cell, txt in zip(t.rows[0].cells, ["Proc./Serv./Doc.", "Resumo", "Deliberações da reunião"]):
        _shd(cell, "D6E4F0"); _brd(cell)
        run = cell.paragraphs[0].add_run(txt)
        run.bold = True; run.font.size = Pt(9); run.font.color.rgb = COR_TITULO
    for ln in linhas:
        row = t.add_row().cells
        for c in row: _brd(c)
        p0 = row[0].paragraphs[0]
        r1 = p0.add_run(ln["id"] + "\n"); r1.bold = True; r1.font.size = Pt(9)
        r2 = p0.add_run(ln.get("desc", "")[:90])
        r2.font.size = Pt(8); r2.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
        _resumo_para(row[1].paragraphs[0], ln.get("resumo", ""))
        row[2].paragraphs[0].add_run("").font.size = Pt(9)
    for row in t.rows:
        for cell, w in zip(row.cells, [Cm(3.5), Cm(9.0), Cm(4.0)]):
            cell.width = w
    doc.add_paragraph()

def tabela_vazia(doc, n=3):
    t = doc.add_table(rows=1 + n, cols=2)
    t.style = "Table Grid"
    for cell, txt in zip(t.rows[0].cells, ["Resumo e status", "Deliberações da reunião"]):
        _shd(cell, "D6E4F0"); _brd(cell)
        run = cell.paragraphs[0].add_run(txt)
        run.bold = True; run.font.size = Pt(9); run.font.color.rgb = COR_TITULO
    for row in t.rows[1:]:
        for c in row.cells: _brd(c)
    doc.add_paragraph()

def gerar_docx(df, model, data_reuniao, participantes, progress_cb=None):
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(2); sec.bottom_margin = Cm(2)
        sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.5)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    total_casos = len(df)
    casos_feitos = 0
    primeira_area = True

    for area in ORDEM_AREAS:
        df_area = df[df["dept"] == area]
        if df_area.empty:
            continue

        if not primeira_area:
            doc.add_page_break()
        primeira_area = False

        # Cabeçalho
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("ATA DE ATENDIMENTO MENSAL")
        run.bold = True; run.font.size = Pt(14); run.font.color.rgb = COR_TITULO

        for label, valor in [("Data:", data_reuniao), ("Participantes:", participantes), (f"Área: {area}", "")]:
            p2 = doc.add_paragraph()
            r1 = p2.add_run(label + " "); r1.bold = True; r1.font.size = Pt(10)
            if valor:
                p2.add_run(valor).font.size = Pt(10)
        doc.add_paragraph()

        # Grupos semelhantes
        grupos    = detectar_grupos(df_area)
        agrupados = {i for g in grupos for i in g}

        # Processos em andamento
        _subtitulo(doc, "Processos em andamento")
        linhas = []; processados = set()

        for grupo in grupos:
            ids_v = [i for i in grupo if i in df_area["id_caso"].values]
            if len(ids_v) < 2 or all(i in processados for i in ids_v):
                continue
            resumo = resumir_grupo(model, ids_v, df_area)
            rows_g = df_area[df_area["id_caso"].isin(ids_v)]
            linhas.append({
                "id":     " | ".join(ids_v),
                "desc":   f"{rows_g.iloc[0]['acao']} — {len(ids_v)} casos semelhantes",
                "resumo": resumo,
            })
            processados.update(ids_v)
            casos_feitos += len(ids_v)
            if progress_cb:
                progress_cb(casos_feitos / total_casos, f"Agrupando {len(ids_v)} casos semelhantes...")

        for _, caso in df_area.iterrows():
            if caso["id_caso"] in processados:
                continue
            resumo = resumir_caso(model, caso)
            linhas.append({"id": caso["id_caso"], "desc": caso["titulo"][:90], "resumo": resumo})
            processados.add(caso["id_caso"])
            casos_feitos += 1
            if progress_cb:
                progress_cb(casos_feitos / total_casos, f"Resumindo {caso['id_caso']}...")

        tabela_casos(doc, linhas)

        # Casos encerrados — vazio para preenchimento manual
        _subtitulo(doc, "Casos encerrados no último mês")
        tabela_vazia(doc, 3)

        # Outras deliberações — vazio para preenchimento manual
        _subtitulo(doc, "Outras solicitações ou deliberações")
        tabela_vazia(doc, 3)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# ─────────────────────────────────────────────────────────────
# INTERFACE STREAMLIT
# ─────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="Gerador de Ata Jurídica",
    page_icon="⚖️",
    layout="centered",
)

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Playfair+Display:wght@600;700&family=Source+Sans+3:wght@300;400;600&display=swap');

html, body, [class*="css"] {
    font-family: 'Source Sans 3', sans-serif;
}
.header-box {
    background: linear-gradient(135deg, #1F497D 0%, #163966 100%);
    border-radius: 12px;
    padding: 32px 36px 28px;
    margin-bottom: 32px;
    position: relative;
    overflow: hidden;
}
.header-box::after {
    content: '⚖';
    position: absolute;
    right: 24px; top: 8px;
    font-size: 90px;
    opacity: .06;
    line-height: 1;
}
.header-box h1 {
    font-family: 'Playfair Display', serif;
    color: white;
    font-size: 28px;
    margin: 0 0 6px;
    font-weight: 700;
}
.header-box p {
    color: rgba(255,255,255,.55);
    font-size: 13px;
    margin: 0;
    text-transform: uppercase;
    letter-spacing: .8px;
    font-weight: 300;
}
.section-label {
    font-size: 11px;
    font-weight: 700;
    color: #1F497D;
    text-transform: uppercase;
    letter-spacing: 1.5px;
    margin-bottom: 10px;
    padding-bottom: 6px;
    border-bottom: 2px solid #D6E4F0;
}
.stat-card {
    background: #F4F6F9;
    border: 1px solid #DDE3EC;
    border-radius: 8px;
    padding: 14px 18px;
    text-align: center;
}
.stat-card .num {
    font-family: 'Playfair Display', serif;
    font-size: 28px;
    color: #1F497D;
    font-weight: 700;
    line-height: 1;
}
.stat-card .lbl {
    font-size: 11px;
    color: #6B7A94;
    text-transform: uppercase;
    letter-spacing: .8px;
    margin-top: 4px;
}
.success-banner {
    background: linear-gradient(135deg, #EBF7EE, #F0F9F2);
    border: 1.5px solid #A8D5B0;
    border-radius: 10px;
    padding: 24px;
    text-align: center;
    margin-top: 20px;
}
.success-banner h3 {
    font-family: 'Playfair Display', serif;
    color: #1E6B35;
    margin: 8px 0 4px;
    font-size: 20px;
}
.success-banner p { color: #4A7A58; font-size: 14px; margin: 0; }
</style>
""", unsafe_allow_html=True)

# Header
st.markdown("""
<div class="header-box">
  <h1>Gerador de Ata Jurídica</h1>
  <p>Família Pilatti &nbsp;·&nbsp; Escritório de Advocacia</p>
</div>
""", unsafe_allow_html=True)

# ── CONFIGURAÇÕES ────────────────────────────────────────────
st.markdown('<div class="section-label">Configurações da Reunião</div>', unsafe_allow_html=True)

col1, col2 = st.columns(2)
with col1:
    data_reuniao = st.text_input("Data da Reunião", placeholder="Ex: 06/03/2026",
                                  value=date.today().strftime("%d/%m/%Y"))
with col2:
    participantes = st.text_input("Participantes", placeholder="Ex: Dr. João, Dra. Ana")

st.markdown("<br>", unsafe_allow_html=True)

# ── UPLOAD ────────────────────────────────────────────────────
st.markdown('<div class="section-label">Planilhas de Casos</div>', unsafe_allow_html=True)

excel1 = st.file_uploader(
    "📊 Excel 1 — Processos judiciais e administrativos",
    type=["xlsx", "xls"],
    help="Qualquer nome de arquivo. Colunas obrigatórias: Pasta - Título, Andamentos / Data/hora, Andamentos / Descrição, Ação, partes"
)

excel2 = st.file_uploader(
    "➕ Excel 2 — Privado / Compliance (opcional)",
    type=["xlsx", "xls"],
    help="Mesmo formato do Excel 1. Adicione uma coluna 'Departamento' com Privado ou Compliance."
)

st.markdown("<br>", unsafe_allow_html=True)

# ── BOTÃO GERAR ───────────────────────────────────────────────
gerar = st.button("⚖️  Gerar Ata", type="primary", use_container_width=True, disabled=not excel1)

if not excel1:
    st.caption("Selecione pelo menos o Excel 1 para continuar.")

# ── PROCESSAMENTO ─────────────────────────────────────────────
if gerar and excel1:
    try:
        # Carrega Excel(s)
        with st.spinner("Lendo planilha(s)..."):
            df1 = carregar_excel(excel1)
            dfs = [df1]
            if excel2:
                df2 = carregar_excel(excel2)
                dfs.append(df2)
            df_total = pd.concat(dfs, ignore_index=True)

        # Estatísticas
        st.markdown("<br>", unsafe_allow_html=True)
        st.markdown('<div class="section-label">Planilha carregada</div>', unsafe_allow_html=True)
        cols = st.columns(len(df_total["dept"].value_counts()) + 1)
        with cols[0]:
            st.markdown(f"""
            <div class="stat-card">
              <div class="num">{len(df_total)}</div>
              <div class="lbl">Total de casos</div>
            </div>""", unsafe_allow_html=True)
        for i, (dept, n) in enumerate(df_total["dept"].value_counts().items(), 1):
            with cols[i]:
                st.markdown(f"""
                <div class="stat-card">
                  <div class="num">{n}</div>
                  <div class="lbl">{dept}</div>
                </div>""", unsafe_allow_html=True)

        st.markdown("<br>", unsafe_allow_html=True)

        # Gera ata com progresso
        st.markdown('<div class="section-label">Gerando Ata</div>', unsafe_allow_html=True)
        prog_bar  = st.progress(0)
        prog_text = st.empty()

        model = get_model()

        def atualizar_progresso(pct, msg):
            prog_bar.progress(min(pct, 1.0))
            prog_text.caption(f"✍ {msg}")

        docx_buf = gerar_docx(df_total, model, data_reuniao, participantes, atualizar_progresso)

        prog_bar.progress(1.0)
        prog_text.empty()

        # Sucesso
        st.markdown("""
        <div class="success-banner">
          <div style="font-size:40px">✅</div>
          <h3>Ata gerada com sucesso!</h3>
          <p>Clique no botão abaixo para baixar o documento Word.</p>
        </div>
        """, unsafe_allow_html=True)

        st.markdown("<br>", unsafe_allow_html=True)
        st.download_button(
            label="⬇  Baixar ata_mensal.docx",
            data=docx_buf,
            file_name=f"ata_{date.today().strftime('%Y_%m')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )

        st.info("💡 Após baixar, preencha manualmente no Word: **Casos encerrados** e **Outras deliberações**.")

    except Exception as e:
        st.error(f"❌ Erro ao processar: {e}")
        st.exception(e)
