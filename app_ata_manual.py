import re
import io
import base64
import pandas as pd
import streamlit as st
from datetime import date
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─────────────────────────────────────────────────────────────
# LOGO (base64 embutida)
# ─────────────────────────────────────────────────────────────
LOGO_B64 = "/9j/4AAQSkZJRgABAQEAYABgAAD/2wBDAAIBAQIBAQICAgICAgICAwUDAwMDAwYEBAMFBwYHBwcGBwcICQsJCAgKCAcHCg0KCgsMDAwMBwkODw0MDgsMDAz/2wBDAQICAgMDAwYDAwYMCAcIDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAwMDAz/wAARCAJUAuMDASIAAhEBAxEB/8QAHwAAAQUBAQEBAQEAAAAAAAAAAAECAwQFBgcICQoL/8QAtRAAAgEDAwIEAwUFBAQAAAF9AQIDAAQRBRIhMUEGE1FhByJxFDKBkaEII0KxwRVS0fAkM2JyggkKFhcYGRolJicoKSo0NTY3ODk6Q0RFRkdISUpTVFVWV1hZWmNkZWZnaGlqc3R1dnd4eXqDhIWGh4iJipKTlJWWl5iZmqKjpKWmp6ipqrKztLW2t7i5usLDxMXGx8jJytLT1NXW19jZ2uHi4+Tl5ufo6erx8vP09fb3+Pn6/8QAHwEAAwEBAQEBAQEBAQAAAAAAAAECAwQFBgcICQoL/8QAtREAAgECBAQDBAcFBAQAAQJ3AAECAxEEBSExBhJBUQdhcRMiMoEIFEKRobHBCSMzUvAVYnLRChYkNOEl8RcYGRomJygpKjU2Nzg5OkNERUZHSElKU1RVVldYWVpjZGVmZ2hpanN0dXZ3eHl6goOEhYaHiImKkpOUlZaXmJmaoqOkpaanqKmqsrO0tba3uLm6wsPExcbHyMnK0tPU1dbX2Nna4uPk5ebn6Onq8vP09fb3+Pn6/9oADAMBAAIRAxEAPwD+f+iiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAr9BP+CYf/AAb3eM/+Cr/wQuvGXw4+LfwotJ9JujZaxoOqzX0ep6SOSTGZUS3dSkijcjqxVsMMhkdV/Puvff8Agm1/wUV8ff8ABMP9qDR/iX4Dud7W+LbWNImkK2mv2LMDJazAdjgFXwSjqrDOMEA/Rv8A4gm/2hf+ip/Bn/v/AKl/8iUf8QTf7Qv/AEVP4M/9/wDUv/kSv6Af2Dv25vAP/BRT9mjQfih8O9R+2aNrCeXc2spAu9Hu1A82zuEBOyWMkZ7MpV1JVlY+x0AfzN/8QTf7Qv8A0VP4M/8Af/Uv/kSj/iCb/aF/6Kn8Gf8Av/qX/wAiV/TJRQB/M3/xBN/tC/8ARU/gz/3/ANS/+RKP+IJv9oX/AKKn8Gf+/wDqX/yJX9MlFAH8zn/EE3+0L/0VP4Nf9/8AUv8A5EqjrP8AwZR/tK2yZsPiP8D7sgdJtR1SAn6YsWr+nGigD+Qv9or/AINeP2yv2erG5vU+G1t48021BZ7jwjqkOoyMB/ctmKXT59FhJr4M8a+Btb+G3ia70TxFo+q6BrNg/l3NhqVpJa3Vs39245AGU+xAr++evFP20v8AgnX8GP8AgoR4HOhfFvwFoniuOKNo7S/kj8nUtNz3t7qPE0XPJCttbHzBhxQB/DdRX6gf8Fsf+Dafx5/wTQtL74g/D+61H4jfBlHLXF2YAdV8MKT8ovUQbXi7faECrnh0jyu78v6ACiiigAooooAKKKKACiivb/2Hv+Ccvxl/4KL/ABDbw78JPBOp+JprZlF/f8W+m6UrdGuLl8Rx8AkKTvbadqsRigDxCu0+Bv7OPxA/ac8YL4f+HXgrxT441pgGNnoemTX0sak43OI1OxfVmwBg5Nf0T/8ABOz/AINiv2V/hV8Fv2ctL8T3HgzwJZ+F/E1z4w1a0hup7fWrm4hn2WbzKPKiZFkuFYoMuSzAf0K/s+fsZfCH9lBb0/Cz4Y+BPh4dSCi9bw1oFpprXe3O3zGijUvjJxnOM0AfyV/8G9v7V/w8/Ya/wCChth4/wDij4gtfDXhGy0TU7Wa8lieX97NCEjUKilubPeivnT/AOJLf8PN/wDgsf8A9Hy/+W1pX/xqv6tKKAP5K/8Ah5v/AMFj/wDo+X/y2tK/+NUf8PN/+Cx//R8v/ltaV/8AGq/rUooA/kr/AOHm/wDwWP8A+j5f/La0r/41X61/sIf8HF3xF/aa/wCCmHgv9mr4d/Cy/wDBeu6zbamLTxPquoWt5Z6gbbT7m6YQJbrv2hoAoLuMEt8pGK/ZivzI/YO/4JafBD9h7/gpj+1T8YvBuh65b+Kfi/4hshfTX+oNcWthbR2KK8VomBsWWWa4c55JlVc4RFAB+m9FFFABRRRQAUUUUAFFFFABRRRQAUUUUAFFFFABRRRQAUUUUAFfDf/BQL/gs/8ACX/gnj+1x8E/gz46F/8A2x8Xb37PfaxDJGtp4ftGmW3S5nVgWbc8kpWNRlo4Jmyuxd/3JX4P/wDByV4D+Cmkftn/AAL1f4o6/faPqnjzw5Z6DpC2Gn/aWdEv7qRoZSWURhvtKMJDkA5GCSKAOL/aS/4OgPD37On7T3jL4b6l+z5c3UXg7W7nRptTi8WRCS6MEhTzDGbTCjjkBmxXv/wm/wCDlX4a/EHxL4a03Vvg94/8Mz+IdYsdG+2rqGl39lYtczyQxySyxT71VmT7uw5J6Guq1X9nb9n34WfDPS/h7afB34f+CfA/gy1GirZXPhWy0y1vY4mLJcIjxqkkkjl2eSMbnLEnJJr8I/+ChniD9jr9hf9onwTH+z74t8SW2heJI9LnvLV9K1Kwg0vUmulinhX+0Y97xbRIFQsoH3Tk5oA/s4ooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooA+fP8AgpP/AMFMfht/wS7/AGbbr4h/EW7kMk0v2HQ9DtW/0/X74jItrc9ABgNJIQVjTk5JRWK/4JP/ALcl7/wUU/4J9+Afizq+jQ6Dq3iSG5t9S063YtBb3ltcywTeWxJbYXiLKGJYKRkkmv4v/wBuD9qXxF+2n+1v48+KXiWa7bUPFmsy3kFvPOZl022UiO3tUJJwkMKxxqM9FGeTmv0H/wCCa/7A/wARP+CmH7Bfjb9nTxJ+0Vp3wf8AhD4ot9Qi8P6Xa+Hzqmse0okvhNcRNOUby0EYO1VLE5JwAD97qKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAooooAKKKKACiiigAr9BP+CYf/AAb3eM/+Cr/wQuvGXw4+LfwotJ9JujZaxoOqzX0ep6SOSTGZUS3dSkijcjqxVsMMhkdV/Puvff8Agm1/wUV8ff8ABMP9qDR/iX4Dud7W+LbWNImkK2mv2LMDJazAdjgFXwSjqrDOMEA/Rv8A4gm/2hf+ip/Bn/v/AKl/8iUf8QTf7Qv/AEVP4M/9/wDUv/kSv6Af2Dv25vAP/BRT9mjQfih8O9R+2aNrCeXc2spAu9Hu1A82zuEBOyWMkZ7MpV1JVlY+x0AfzN/8QTf7Qv8A0VP4M/8Af/Uv/kSj/iCb/aF/6Kn8Gf8Av/qX/wAiV/TJRQB/M3/xBN/tC/8ARU/gz/3/ANS/+RKP+IJv9oX/AKKn8Gf+/wDqX/yJX9MlFAH8zn/EE3+0L/0VP4Nf9/8AUv8A5EqjrP8AwZR/tK2yZsPiP8D7sgdJtR1SAn6YsWr+nGigD+Qv9or/AINeP2yv2erG5vU+G1t48021BZ7jwjqkOoyMB/ctmKXT59FhJr4M8a+Btb+G3ia70TxFo+q6BrNg/l3NhqVpJa3Vs39245AGU+xAr++evFP20v8AgnX8GP8AgoR4HOhfFvwFoniuOKNo7S/kj8nUtNz3t7qPE0XPJCttbHzBhxQB/DdRX6gf8Fsf+Dafx5/wTQtL74g/D+61H4jfBlHLXF2YAdV8MKT8ovUQbXi7faECrnh0jyu78v6ACiiigAooooAKKKKA"

ORDEM_AREAS = ["Cível", "Público", "Trabalhista", "Privado", "Compliance"]

COR_TITULO = RGBColor(0x00, 0x00, 0x00)
COR_SUBTIT = RGBColor(0x00, 0x00, 0x00)

MAPA_ACAO_DEPT = {
    "reclamação trabalhista":              "Trabalhista",
    "reclamatória trabalhista":            "Trabalhista",
    "processo administrativo trabalhista": "Trabalhista",
    "mandado de segurança":                "Público",
    "execução fiscal":                     "Público",
    "processo administrativo":             "Público",
    "auditoria fiscal":                    "Público",
    "ação ordinária":                      "Público",
    "reclamação procon":                   "Cível",
    "notificação":                         "Cível",
    "indenizatória":                       "Cível",
    "ajuizamento pendente":                "Cível",
    "usucapião":                           "Cível",
    "execução de título extrajudicial":    "Cível",
    "demarcatória":                        "Cível",
    "despejo":                             "Cível",
    "falência":                            "Cível",
}

# ─────────────────────────────────────────────────────────────
# PARSING DE TEXTO COLADO
# ─────────────────────────────────────────────────────────────
def inferir_dept(acao, orgao=""):
    a = acao.lower().strip()
    o = orgao.lower().strip()
    for chave, dept in MAPA_ACAO_DEPT.items():
        if chave in a:
            return dept
    if any(k in o for k in ["trabalho", "trt", "ministério público do trabalho"]):
        return "Trabalhista"
    if any(k in o for k in ["fiscal", "receita", "fazenda", "tribunal administrativo"]):
        return "Público"
    return "Cível"


def parsear_texto_livre(texto_bruto: str) -> pd.DataFrame:
    """
    Interpreta texto colado pelo usuário.

    Formato esperado (blocos separados por linha em branco ou '---'):
    Pasta: Proc-0001
    Ação: Reclamação Trabalhista
    Partes: Empresa X x Fulano
    Órgão: 1ª Vara do Trabalho
    Valor: R$ 50.000
    Distribuído em: 01/01/2023
    Área: Trabalhista          ← opcional
    Andamentos:
    [15/02/2026] Audiência realizada. Acordo frustrado.
    [10/01/2026] Contestação protocolada.
    ---
    Pasta: Proc-0002
    ...
    """
    casos = []
    # Divide em blocos por linha "---" ou dupla linha em branco
    blocos = re.split(r"\n---+\n|\n{3,}", texto_bruto.strip())

    for bloco in blocos:
        bloco = bloco.strip()
        if not bloco:
            continue

        caso = {
            "id_caso": "", "titulo": "", "dept": "", "acao": "",
            "partes": "", "orgao": "", "valor": "", "historico": "",
            "ultimo": "", "n_and": 0, "data_distribuicao": "",
            "resumo_manual": "",
        }

        # Extrai andamentos (tudo após "Andamentos:")
        partes_and = re.split(r"(?im)^andamentos\s*:\s*\n", bloco, maxsplit=1)
        cabecalho = partes_and[0]
        historico_raw = partes_and[1].strip() if len(partes_and) > 1 else ""

        # Extrai resumo manual (após "Resumo:" se existir)
        partes_res = re.split(r"(?im)^resumo\s*:\s*\n?", cabecalho, maxsplit=1)
        cabecalho = partes_res[0]
        if len(partes_res) > 1:
            caso["resumo_manual"] = partes_res[1].strip()

        # Parseia campos do cabeçalho
        for linha in cabecalho.splitlines():
            linha = linha.strip()
            if not linha:
                continue
            for campo, chave in [
                (r"(?i)^pasta\s*:", "id_caso"),
                (r"(?i)^a[çc][ãa]o\s*:", "acao"),
                (r"(?i)^partes?\s*:", "partes"),
                (r"(?i)^[oó]rg[aã]o\s*:", "orgao"),
                (r"(?i)^valor\s*:", "valor"),
                (r"(?i)^distribu[ií]d[oa]\s+em\s*:", "data_distribuicao"),
                (r"(?i)^[áa]rea\s*:", "dept"),
            ]:
                m = re.match(campo + r"\s*(.*)", linha)
                if m:
                    caso[chave] = m.group(1).strip()
                    break

        # Fallback: usa primeira linha como id_caso se não encontrou
        if not caso["id_caso"]:
            primeiras = [l.strip() for l in cabecalho.splitlines() if l.strip()]
            caso["id_caso"] = primeiras[0] if primeiras else f"Caso-{len(casos)+1}"

        # Historico
        caso["historico"] = historico_raw
        andamentos_lista = [l.strip() for l in historico_raw.splitlines() if l.strip()]
        caso["n_and"] = len(andamentos_lista)
        caso["ultimo"] = andamentos_lista[0] if andamentos_lista else ""

        # Título
        partes_str = caso["partes"] or caso["id_caso"]
        caso["titulo"] = f"{caso['acao']} — {partes_str[:60]}" if caso["acao"] else partes_str[:80]

        # Área
        if not caso["dept"]:
            caso["dept"] = inferir_dept(caso["acao"], caso["orgao"])

        casos.append(caso)

    return pd.DataFrame(casos) if casos else pd.DataFrame()


# ─────────────────────────────────────────────────────────────
# GERAÇÃO DO WORD
# ─────────────────────────────────────────────────────────────
def _shd(cell, cor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    s = OxmlElement("w:shd")
    s.set(qn("w:val"), "clear")
    s.set(qn("w:color"), "auto")
    s.set(qn("w:fill"), cor)
    tcPr.append(s)


def _brd(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcB = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "C8C8C8")
        tcB.append(b)
    tcPr.append(tcB)


def _resumo_para(paragraph, texto):
    paragraph.clear()
    run = paragraph.add_run(texto)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)


def _subtitulo(doc, texto):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(texto)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = COR_SUBTIT


def tabela_casos(doc, linhas):
    if not linhas:
        return
    t = doc.add_table(rows=1, cols=3)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for cell, txt in zip(t.rows[0].cells, ["Proc./Serv./Doc.", "Resumo", "Deliberações da reunião"]):
        _shd(cell, "E8E8E8")
        _brd(cell)
        run = cell.paragraphs[0].add_run(txt)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = COR_TITULO
    for ln in linhas:
        row = t.add_row().cells
        for c in row:
            _brd(c)
        p0 = row[0].paragraphs[0]
        r1 = p0.add_run(ln["id"] + "\n")
        r1.bold = True
        r1.font.size = Pt(9)
        r2 = p0.add_run(ln.get("desc", "")[:90])
        r2.font.size = Pt(8)
        r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        _resumo_para(row[1].paragraphs[0], ln.get("resumo", ""))
        row[2].paragraphs[0].add_run("").font.size = Pt(9)
    for row in t.rows:
        for cell, w in zip(row.cells, [Cm(3.5), Cm(9.0), Cm(4.0)]):
            cell.width = w
    doc.add_paragraph()


def tabela_vazia(doc, n=3):
    t = doc.add_table(rows=1 + n, cols=2)
    t.style = "Table Grid"
    for cell, txt in zip(t.rows[0].cells, ["Resumo e status", "Deliberações da reunião"]):
        _shd(cell, "E8E8E8")
        _brd(cell)
        run = cell.paragraphs[0].add_run(txt)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = COR_TITULO
    for row in t.rows[1:]:
        for c in row.cells:
            _brd(c)
    doc.add_paragraph()


def gerar_docx(df, data_reuniao, participantes):
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(2)
        sec.bottom_margin = Cm(2)
        sec.left_margin = Cm(2.5)
        sec.right_margin = Cm(2.5)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    primeira_area = True

    for area in ORDEM_AREAS:
        df_area = df[df["dept"] == area]
        if df_area.empty:
            continue

        if not primeira_area:
            doc.add_page_break()
        primeira_area = False

        # Cabeçalho
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("ATA DE ATENDIMENTO MENSAL")
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = COR_TITULO

        for label, valor in [
            ("Data:", data_reuniao),
            ("Participantes:", participantes),
            (f"Área: {area}", ""),
        ]:
            p2 = doc.add_paragraph()
            r1 = p2.add_run(label + " ")
            r1.bold = True
            r1.font.size = Pt(10)
            if valor:
                p2.add_run(valor).font.size = Pt(10)
        doc.add_paragraph()

        # Processos em andamento
        _subtitulo(doc, "Processos em andamento")
        linhas = []

        for _, caso in df_area.iterrows():
            # Usa resumo_manual se existir; caso contrário monta automaticamente
            resumo = str(caso.get("resumo_manual", "")).strip()
            if not resumo:
                partes = str(caso.get("partes", "")).strip()
                acao = str(caso.get("acao", "")).strip()
                orgao = str(caso.get("orgao", "")).strip()
                valor = str(caso.get("valor", "")).strip()
                dist = str(caso.get("data_distribuicao", "")).strip()
                ultimo = str(caso.get("ultimo", "")).strip()

                partes_txt = f"Partes: {partes}. " if partes else ""
                orgao_txt = f"Órgão: {orgao}. " if orgao else ""
                valor_txt = f"Valor da causa: {valor}. " if valor and valor not in ("", "nan") else ""
                dist_txt = f"Distribuído em: {dist}. " if dist and dist not in ("", "nan", "NaT") else ""
                ultimo_txt = f"Último andamento: {ultimo[:250]}." if ultimo else "Sem andamentos registrados."

                resumo = f"{acao}. {partes_txt}{orgao_txt}{valor_txt}{dist_txt}{ultimo_txt} Deliberação:".strip()

            linhas.append({
                "id": str(caso["id_caso"]),
                "desc": str(caso.get("titulo", ""))[:90],
                "resumo": resumo,
            })

        tabela_casos(doc, linhas)

        # Casos encerrados
        _subtitulo(doc, "Casos encerrados no último mês")
        tabela_vazia(doc, 3)

        # Outras deliberações
        _subtitulo(doc, "Outras solicitações ou deliberações")
        tabela_vazia(doc, 3)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ─────────────────────────────────────────────────────────────
# INTERFACE STREAMLIT
# ─────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="GERADOR DE ATA DE ATENDIMENTO BCA",
    page_icon="⚖️",
    layout="centered",
)

st.markdown("""
<style>
*, html, body, [class*="css"] {
    font-family: 'Helvetica Neue', Helvetica, Arial, sans-serif !important;
}
#MainMenu, footer, header { visibility: hidden; }
.stApp { background-color: #F0F0F0; }
.block-container { padding-top: 2rem !important; max-width: 800px !important; }

.bca-header {
    background: #1A1A1A;
    border-radius: 10px;
    padding: 0;
    margin-bottom: 28px;
    overflow: hidden;
    display: flex;
    align-items: stretch;
    box-shadow: 0 4px 24px rgba(0,0,0,0.18);
}
.bca-logo-wrap {
    background: #111;
    padding: 20px 28px;
    display: flex;
    align-items: center;
    justify-content: center;
    border-right: 1px solid #2a2a2a;
    min-width: 140px;
}
.bca-logo-wrap img { width: 110px; display: block; }
.bca-title-wrap {
    padding: 20px 28px;
    display: flex;
    flex-direction: column;
    justify-content: center;
}
.bca-title-wrap h1 {
    color: #FFFFFF;
    font-size: 16px;
    font-weight: 700;
    margin: 0 0 5px;
    letter-spacing: .4px;
    line-height: 1.3;
    text-transform: uppercase;
}
.bca-title-wrap p {
    color: #888;
    font-size: 11px;
    margin: 0;
    text-transform: uppercase;
    letter-spacing: 1.2px;
}

.sec-label {
    font-size: 10px;
    font-weight: 700;
    color: #333;
    text-transform: uppercase;
    letter-spacing: 1.8px;
    margin-bottom: 10px;
    padding-bottom: 6px;
    border-bottom: 1.5px solid #D0D0D0;
}

.stat-card {
    background: #fff;
    border: 1px solid #D8D8D8;
    border-radius: 7px;
    padding: 14px 18px;
    text-align: center;
    box-shadow: 0 1px 4px rgba(0,0,0,.06);
}
.stat-card .num { font-size: 26px; color: #1A1A1A; font-weight: 700; line-height: 1; }
.stat-card .lbl { font-size: 10px; color: #555; text-transform: uppercase; letter-spacing: .8px; margin-top: 4px; }

.success-box {
    background: #F6F6F6;
    border: 1.5px solid #C0C0C0;
    border-radius: 9px;
    padding: 24px;
    text-align: center;
    margin-top: 16px;
}
.success-box h3 { color: #1A1A1A; font-size: 18px; margin: 8px 0 4px; }
.success-box p  { color: #444; font-size: 13px; margin: 0; }

.stButton > button {
    background: #1A1A1A !important;
    color: #fff !important;
    border: none !important;
    border-radius: 7px !important;
    font-size: 14px !important;
    font-weight: 600 !important;
    padding: 14px !important;
    transition: opacity .2s !important;
}
.stButton > button:hover { opacity: .85 !important; }

.stDownloadButton > button {
    background: #333 !important;
    color: #fff !important;
    border-radius: 7px !important;
    font-weight: 600 !important;
    width: 100% !important;
}

/* Inputs com texto legível */
label, .stTextInput label, .stTextArea label, .stRadio label,
div[data-testid="stWidgetLabel"] > p,
div[data-testid="stWidgetLabel"] > label {
    color: #111111 !important;
    font-weight: 600 !important;
    font-size: 13px !important;
}

.stTextInput > div > div > input,
.stTextArea > div > textarea {
    border: 1.5px solid #BBBBBB !important;
    border-radius: 6px !important;
    background: #fff !important;
    color: #111111 !important;
    font-size: 13px !important;
}
.stTextInput > div > div > input::placeholder,
.stTextArea > div > textarea::placeholder {
    color: #888888 !important;
}
.stTextInput > div > div > input:focus,
.stTextArea > div > textarea:focus {
    border-color: #444 !important;
    box-shadow: 0 0 0 2px rgba(0,0,0,.08) !important;
}

/* Radio buttons */
.stRadio > div { gap: 16px !important; }
.stRadio > div label { color: #111111 !important; font-size: 13px !important; }

/* File uploader */
[data-testid="stFileUploader"] {
    background: #fff;
    border: 1.5px dashed #C0C0C0;
    border-radius: 8px;
    padding: 4px;
}
[data-testid="stFileUploader"] label,
[data-testid="stFileUploader"] p,
[data-testid="stFileUploader"] span {
    color: #222222 !important;
}

.stProgress > div > div > div { background: #1A1A1A !important; }
.stAlert { border-radius: 7px !important; }

/* Tabs */
.stTabs [data-baseweb="tab"] {
    color: #333 !important;
    font-weight: 600 !important;
    font-size: 13px !important;
}
.stTabs [aria-selected="true"] {
    color: #000 !important;
    border-bottom: 2px solid #1A1A1A !important;
}

/* Caption */
.stCaption, [data-testid="stCaptionContainer"] p {
    color: #555 !important;
    font-size: 12px !important;
}

/* Info/warning boxes */
[data-testid="stAlert"] p { color: #111 !important; }

</style>
""", unsafe_allow_html=True)

# ── HEADER ───────────────────────────────────────────────────
st.markdown(f"""
<div class="bca-header">
  <div class="bca-logo-wrap">
    <img src="data:image/jpeg;base64,{LOGO_B64}" alt="BCA">
  </div>
  <div class="bca-title-wrap">
    <h1>Gerador de Ata de<br>Atendimento BCA</h1>
    <p>Controle de Operações</p>
  </div>
</div>
""", unsafe_allow_html=True)

# ── CONFIGURAÇÕES DA REUNIÃO ──────────────────────────────────
st.markdown('<div class="sec-label">Configurações da Reunião</div>', unsafe_allow_html=True)

col1, col2 = st.columns(2)
with col1:
    data_reuniao = st.text_input(
        "Data da Reunião",
        placeholder="Ex: 06/03/2026",
        value=date.today().strftime("%d/%m/%Y"),
    )
with col2:
    participantes = st.text_input(
        "Participantes",
        placeholder="Ex: Dr. João, Dra. Ana",
    )

st.markdown("<br>", unsafe_allow_html=True)

# ── ABAS DE ENTRADA ───────────────────────────────────────────
st.markdown('<div class="sec-label">Dados dos Processos</div>', unsafe_allow_html=True)

aba_texto, aba_pdf, aba_excel = st.tabs([
    "✍️  Colar Texto",
    "📄  Upload PDF",
    "📊  Upload Excel",
])

df_total = None
fonte = None

# ── ABA 1: TEXTO LIVRE ────────────────────────────────────────
with aba_texto:
    st.markdown("""
    <p style="color:#333;font-size:13px;margin-bottom:8px;">
    Cole os dados dos processos abaixo.<br>
    Separe cada processo com uma linha <code>---</code><br>
    Use o formato abaixo. O campo <strong>Resumo:</strong> é opcional — se ausente, o app monta automaticamente.
    </p>
    """, unsafe_allow_html=True)

    EXEMPLO = """Pasta: Proc-0001
Ação: Reclamação Trabalhista
Partes: Empresa ABC Ltda x João da Silva
Órgão: 1ª Vara do Trabalho de São Paulo
Valor: R$ 45.000,00
Distribuído em: 10/03/2023
Área: Trabalhista
Resumo: Reclamação trabalhista movida por ex-funcionário pleiteando verbas rescisórias. Audiência realizada em fevereiro sem acordo. Processo em fase de instrução com perícia agendada para abril/2026. Deliberação:
Andamentos:
[15/02/2026] Audiência de instrução realizada. Sem acordo entre as partes.
[10/01/2026] Perícia contábil deferida pelo juízo.
[05/12/2025] Contestação protocolada com impugnação à inicial.
---
Pasta: Proc-0002
Ação: Execução Fiscal
Partes: Município de São Paulo x Empresa ABC Ltda
Órgão: 3ª Vara da Fazenda Pública
Valor: R$ 120.000,00
Distribuído em: 05/06/2021
Área: Público
Andamentos:
[20/02/2026] Penhora de bem imóvel determinada pelo juízo.
[15/01/2026] Garantia do juízo apresentada. Aguarda decisão sobre embargos."""

    texto_colado = st.text_area(
        "Dados dos processos",
        height=320,
        placeholder=EXEMPLO,
        label_visibility="collapsed",
    )

    if texto_colado.strip():
        with st.expander("📋 Ver formato esperado / exemplo", expanded=False):
            st.code(EXEMPLO, language="text")
        df_texto = parsear_texto_livre(texto_colado)
        if not df_texto.empty:
            fonte = "texto"
            df_total = df_texto
            st.success(f"✅ {len(df_texto)} processo(s) reconhecido(s)")
        else:
            st.warning("⚠️ Não foi possível reconhecer processos. Verifique o formato.")
    else:
        with st.expander("📋 Ver formato esperado / exemplo", expanded=True):
            st.code(EXEMPLO, language="text")

# ── ABA 2: PDF ────────────────────────────────────────────────
with aba_pdf:
    st.markdown("""
    <p style="color:#333;font-size:13px;margin-bottom:8px;">
    Faça upload do PDF exportado do <strong>Novajus</strong> — relatório de Andamentos de Casos.
    </p>
    """, unsafe_allow_html=True)

    arquivo_pdf = st.file_uploader(
        "PDF exportado do Novajus",
        type=["pdf"],
        key="uploader_pdf",
        label_visibility="collapsed",
    )

    if arquivo_pdf:
        try:
            import pdfplumber

            def parsear_pdf(uploaded_file):
                CAMPOS = {
                    "Número CNJ": "cnj", "Ação": "acao", "Natureza": "natureza",
                    "Data da distribuição": "data_distribuicao", "Valor da causa": "valor",
                    "Status": "status", "Escritório responsável": "escritorio",
                    "Cliente principal": "cliente", "Contrário principal": "contrario",
                    "Órgão": "orgao",
                }

                def cel(row, i):
                    return (row[i] or "").strip() if i < len(row) else ""

                casos_pdf = []
                caso_pdf = None

                def salvar_pdf():
                    if not caso_pdf or not caso_pdf.get("acao"):
                        return
                    ands = caso_pdf.get("_ands", [])
                    historico = "\n".join(f"[{a['data']}] {a['desc']}" for a in ands)
                    ultimo = ands[0]["desc"] if ands else ""
                    esc = caso_pdf.get("escritorio", "").lower()
                    if "trabalhista" in esc:
                        dept = "Trabalhista"
                    elif "público" in esc or "publico" in esc:
                        dept = "Público"
                    elif "cível" in esc or "conflitos" in esc:
                        dept = "Cível"
                    elif "privado" in esc:
                        dept = "Privado"
                    elif "compliance" in esc:
                        dept = "Compliance"
                    else:
                        dept = inferir_dept(caso_pdf.get("acao", ""))
                    cliente = caso_pdf.get("cliente", "").replace("(filial cliente principal)", "").strip()
                    contrario = caso_pdf.get("contrario", "")
                    casos_pdf.append({
                        "id_caso": caso_pdf["_pasta"],
                        "titulo": f"{caso_pdf.get('acao','')} — {cliente[:40]} x {contrario[:40]}",
                        "dept": dept, "acao": caso_pdf.get("acao", ""),
                        "partes": f"{cliente} x {contrario}",
                        "orgao": caso_pdf.get("orgao", ""), "valor": caso_pdf.get("valor", ""),
                        "historico": historico, "ultimo": ultimo, "n_and": len(ands),
                        "data_distribuicao": caso_pdf.get("data_distribuicao", ""),
                        "resumo_manual": "",
                    })

                with pdfplumber.open(uploaded_file) as pdf:
                    for page in pdf.pages:
                        for tab in page.extract_tables():
                            for row in tab:
                                if cel(row, 1) == "Pasta" and cel(row, 2).startswith("Proc"):
                                    salvar_pdf()
                                    caso_pdf = {"_pasta": cel(row, 2).replace(" ", ""), "_ands": []}
                                    continue
                                if caso_pdf is None:
                                    continue
                                campo = cel(row, 2)
                                if campo in CAMPOS:
                                    caso_pdf[CAMPOS[campo]] = cel(row, 6)
                                    continue
                                data_c = cel(row, 3)
                                tipo_c = cel(row, 5)
                                desc_c = cel(row, 10)
                                if re.match(r"\d{2}/\d{2}/\d{4}", data_c) and "Andamento" in tipo_c and desc_c:
                                    caso_pdf["_ands"].append({"data": data_c, "desc": " ".join(desc_c.split())})
                salvar_pdf()
                return pd.DataFrame(casos_pdf)

            df_pdf = parsear_pdf(arquivo_pdf)
            if not df_pdf.empty:
                fonte = "pdf"
                df_total = df_pdf
                st.success(f"✅ {len(df_pdf)} processo(s) lido(s) do PDF")
            else:
                st.warning("⚠️ Nenhum processo encontrado no PDF.")
        except ImportError:
            st.error("❌ Instale pdfplumber: `pip install pdfplumber`")
        except Exception as e:
            st.error(f"❌ Erro ao ler PDF: {e}")

# ── ABA 3: EXCEL ──────────────────────────────────────────────
with aba_excel:
    st.markdown("""
    <p style="color:#333;font-size:13px;margin-bottom:8px;">
    Faça upload da planilha Excel com os processos.
    </p>
    """, unsafe_allow_html=True)

    arquivo_xlsx = st.file_uploader(
        "Excel — Processos",
        type=["xlsx", "xls"],
        key="uploader_xlsx",
        label_visibility="collapsed",
    )
    arquivo_xlsx2 = st.file_uploader(
        "Excel 2 — Privado / Compliance (opcional)",
        type=["xlsx", "xls"],
        key="uploader_xlsx2",
    )

    if arquivo_xlsx:
        try:
            def carregar_excel(f):
                df_raw = pd.read_excel(f, dtype=str).fillna("")
                df_raw.columns = [str(c).strip().lower().replace(" ", "_") for c in df_raw.columns]
                mapa = {
                    "pasta": "pasta", "n°_do_processo": "pasta", "proc": "pasta",
                    "descricao": "descricao", "andamento": "descricao", "histórico": "descricao",
                    "data": "data_str", "data_do_andamento": "data_str",
                    "responsavel": "responsavel", "advogado": "responsavel",
                    "cnj": "cnj", "número_cnj": "cnj",
                    "valor": "valor", "valor_da_causa": "valor",
                    "órgão": "orgao", "orgao": "orgao", "tribunal": "orgao",
                    "ação": "acao", "acao": "acao", "tipo_de_ação": "acao",
                    "partes": "partes", "parte": "partes",
                    "área": "dept_col", "area": "dept_col", "departamento": "dept_col",
                    "data_de_distribuição": "data_distribuicao",
                }
                df_raw = df_raw.rename(columns={c: mapa[c] for c in df_raw.columns if c in mapa})
                for col in ["responsavel", "cnj", "valor", "orgao", "dept_col", "data_distribuicao",
                            "acao", "partes"]:
                    if col not in df_raw.columns:
                        df_raw[col] = ""
                if "pasta" not in df_raw.columns:
                    st.error("Coluna 'Pasta' não encontrada na planilha.")
                    return pd.DataFrame()
                if "descricao" not in df_raw.columns:
                    df_raw["descricao"] = ""
                if "data_str" not in df_raw.columns:
                    df_raw["data_str"] = ""
                df_raw["data_parsed"] = pd.to_datetime(df_raw["data_str"], errors="coerce")
                casos_xl = []
                for pasta, grupo in df_raw.groupby("pasta", sort=False):
                    primeira = grupo.iloc[0]
                    historico = "\n".join(
                        f"[{str(r['data_parsed'])[:10]}] {str(r['descricao']).strip()}"
                        for _, r in grupo.iterrows() if str(r["descricao"]).strip()
                    )
                    ultimo = grupo.iloc[-1]["descricao"].strip()[:300] if len(grupo) else ""
                    dept = primeira["dept_col"].strip() if primeira["dept_col"].strip() \
                        else inferir_dept(str(primeira["acao"]), str(primeira["orgao"]))
                    casos_xl.append({
                        "id_caso": str(pasta).strip(),
                        "titulo": f"{str(primeira['acao']).strip()} — {str(primeira['partes']).strip()[:60]}",
                        "dept": dept, "acao": str(primeira["acao"]).strip(),
                        "partes": str(primeira["partes"]).strip(),
                        "orgao": str(primeira["orgao"]).strip(),
                        "valor": str(primeira["valor"]).strip(),
                        "historico": historico, "ultimo": ultimo, "n_and": len(grupo),
                        "data_distribuicao": str(primeira["data_distribuicao"]).strip(),
                        "resumo_manual": "",
                    })
                return pd.DataFrame(casos_xl)

            dfs = [carregar_excel(arquivo_xlsx)]
            if arquivo_xlsx2:
                dfs.append(carregar_excel(arquivo_xlsx2))
            df_xl = pd.concat(dfs, ignore_index=True)
            if not df_xl.empty:
                fonte = "excel"
                df_total = df_xl
                st.success(f"✅ {len(df_xl)} processo(s) lido(s) do Excel")
        except Exception as e:
            st.error(f"❌ Erro ao ler Excel: {e}")

st.markdown("<br>", unsafe_allow_html=True)

# ── ESTATÍSTICAS ──────────────────────────────────────────────
if df_total is not None and not df_total.empty:
    st.markdown('<div class="sec-label">Processos carregados</div>', unsafe_allow_html=True)
    vcounts = df_total["dept"].value_counts()
    cols = st.columns(len(vcounts) + 1)
    with cols[0]:
        st.markdown(f'''<div class="stat-card">
          <div class="num">{len(df_total)}</div>
          <div class="lbl">Total</div>
        </div>''', unsafe_allow_html=True)
    for i, (dept, n) in enumerate(vcounts.items(), 1):
        with cols[i]:
            st.markdown(f'''<div class="stat-card">
              <div class="num">{n}</div>
              <div class="lbl">{dept}</div>
            </div>''', unsafe_allow_html=True)
    st.markdown("<br>", unsafe_allow_html=True)

# ── BOTÃO GERAR ───────────────────────────────────────────────
pode_gerar = df_total is not None and not df_total.empty
gerar = st.button(
    "⚖️  Gerar Ata",
    type="primary",
    use_container_width=True,
    disabled=not pode_gerar,
)
if not pode_gerar:
    st.caption("Cole os dados ou faça upload de um arquivo para continuar.")

# ── GERAÇÃO ───────────────────────────────────────────────────
if gerar and pode_gerar:
    try:
        with st.spinner("Gerando documento Word..."):
            docx_buf = gerar_docx(df_total, data_reuniao, participantes)

        st.markdown("""
        <div class="success-box">
          <div style="font-size:36px">✅</div>
          <h3>Ata gerada com sucesso!</h3>
          <p>Clique abaixo para baixar o documento Word.</p>
        </div>""", unsafe_allow_html=True)

        st.markdown("<br>", unsafe_allow_html=True)
        st.download_button(
            label="⬇  Baixar Ata de Atendimento BCA.docx",
            data=docx_buf,
            file_name=f"ATA_BCA_{date.today().strftime('%Y_%m')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )
        st.info("💡 Após baixar, preencha manualmente no Word: **Casos encerrados** e **Outras deliberações**.")

    except Exception as e:
        st.error(f"❌ Erro ao gerar o Word: {type(e).__name__}: {e}")
        st.exception(e)
